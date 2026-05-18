import unittest

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from thesis_config import dump_default_config
from thesis_formatter.headers import setup_headers
from thesis_formatter.page import setup_page_numbers


def _make_config():
    cfg = yaml.safe_load(dump_default_config())
    cfg["header_footer"]["enabled"] = True
    cfg["header_footer"]["scope"] = "body"
    cfg["header_footer"]["odd_page_text"] = "ODD_HEADER"
    cfg["header_footer"]["even_page_text"] = "EVEN_HEADER"
    cfg["header_footer"]["different_odd_even"] = True
    cfg["header_footer"]["border_bottom"] = False
    cfg["page_numbers"]["front_position"] = "center"
    cfg["page_numbers"]["body_position"] = "alternate"
    cfg["page_numbers"]["body_odd_position"] = "right"
    cfg["page_numbers"]["body_even_position"] = "left"
    return cfg


def _heading_1(doc, text):
    para = doc.add_paragraph(text)
    para.style = doc.styles["Heading 1"]
    return para


def _build_doc_with_front_matter_and_two_body_sections():
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是前置部分。")
    doc.add_paragraph("关键词：测试；页眉；页码")

    _heading_1(doc, "第1章 绪论")
    doc.add_paragraph("这是正文第一节。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    _heading_1(doc, "第2章 方法")
    doc.add_paragraph("这是正文第二节。")
    return doc


def _build_doc_with_two_front_sections_before_body():
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是中文摘要内容。")
    doc.add_paragraph("关键词：测试；分页；页码")

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 绪论 ................................ 1")
    doc.add_paragraph("第二章 方法 ................................ 5")

    _heading_1(doc, "第1章 绪论")
    doc.add_paragraph("这是正文第一页。")
    return doc


def _section_pg_num_attrs(section):
    node = section._sectPr.find(qn("w:pgNumType"))
    if node is None:
        return {}
    return {key.split("}", 1)[-1]: value for key, value in node.attrib.items()}


class MultiSectionBodyHeaderTests(unittest.TestCase):
    def test_body_headers_and_page_number_alignment_apply_to_all_body_sections(self):
        doc = _build_doc_with_front_matter_and_two_body_sections()
        cfg = _make_config()

        setup_page_numbers(doc, cfg)
        setup_headers(doc, cfg)

        self.assertEqual(len(doc.sections), 3)

        front, body_first, body_second = doc.sections

        self.assertEqual(front.header.paragraphs[0].text, "")
        self.assertEqual(front.even_page_header.paragraphs[0].text, "")
        self.assertEqual(front.footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

        self.assertEqual(body_first.header.paragraphs[0].text, "ODD_HEADER")
        self.assertEqual(body_first.even_page_header.paragraphs[0].text, "EVEN_HEADER")
        self.assertEqual(body_first.footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(body_first.even_page_footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)

        self.assertEqual(body_second.header.paragraphs[0].text, "ODD_HEADER")
        self.assertEqual(body_second.even_page_header.paragraphs[0].text, "EVEN_HEADER")
        self.assertEqual(body_second.footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(body_second.even_page_footer.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def test_all_front_sections_keep_front_page_number_format_until_body_section(self):
        doc = _build_doc_with_two_front_sections_before_body()
        cfg = _make_config()

        setup_page_numbers(doc, cfg)

        self.assertEqual(len(doc.sections), 3)

        front_first, front_second, body_first = doc.sections

        self.assertEqual(
            _section_pg_num_attrs(front_first),
            {"fmt": "upperRoman", "start": "1"},
        )
        self.assertEqual(
            _section_pg_num_attrs(front_second),
            {"fmt": "upperRoman"},
        )
        self.assertEqual(
            _section_pg_num_attrs(body_first),
            {"fmt": "decimal", "start": "1"},
        )


if __name__ == "__main__":
    unittest.main()
