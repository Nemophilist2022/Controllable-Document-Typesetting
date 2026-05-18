"""Hand-built docx fixtures for the MVP end-to-end tests.

Strict rule (T13): these fixtures must NOT be produced by the
``thesis_runner.run_format`` / ``thesis_format_cli.py`` fast path. We
build them straight with python-docx so the "perfect" fixture only
satisfies the rules in the SCAU profile — nothing more, nothing less.

A2 expanded the SCAU profile from 4 to 41 rules, so this builder now
sets:
- page margins / gutter / header & footer distance (SCAU values)
- Normal-style font / size / line spacing / first-line indent
- one Heading 1 paragraph
- one TOC 1 paragraph (matches the 1 heading)
- the four front-matter present-checks: 摘要 / 关键词 / Abstract / Key words
"""

from __future__ import annotations

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_east_asia_font(style, name: str) -> None:
    """Set the eastAsia attribute on the style's rPr/rFonts element."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _ensure_toc1_style(doc):
    """python-docx's default style table doesn't include 'TOC 1', so we
    add it on the fly. Returns the style object."""
    if "TOC 1" in [s.name for s in doc.styles]:
        return doc.styles["TOC 1"]
    return doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)


def _apply_scau_page_layout(doc) -> None:
    """Apply the page-layout values from defaults/scau_2024.yaml so the
    page.* rules pass. We touch every existing section (a fresh docx
    has exactly one) — the values are intentionally hard-coded here so
    fixture invariants are not coupled to YAML changes."""
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
        section.gutter = Cm(0.5)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


# ---------------------------------------------------------------------------
# Public builders
# ---------------------------------------------------------------------------

def build_perfect_docx(path: str) -> None:
    """A minimal docx where all rules in the SCAU profile pass.

    Intentionally NOT formatted by the fast path; we set only what the
    rules check, nothing more. Tools added in A1 (sections, captions,
    references...) are exercised via the orchestrator's plan but they
    don't have to *change* anything for the fixture to pass.
    """
    doc = Document()

    _apply_scau_page_layout(doc)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    _set_east_asia_font(normal, "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(24)

    _ensure_toc1_style(doc)

    # Front-matter — the four ``front_matter.*.present`` rules want
    # these labels somewhere in the document. The text content itself
    # is irrelevant; our checks only look for the title / prefix.
    doc.add_paragraph("摘要", style="Normal")
    doc.add_paragraph("关键词：示例; 论文; 排版", style="Normal")
    doc.add_paragraph("Abstract: This is a placeholder abstract.", style="Normal")
    doc.add_paragraph("Key words: thesis; formatting; rules", style="Normal")

    # heading.h1.style_present: at least one Heading 1 paragraph.
    # NOTE: text intentionally avoids the "第N章" pattern so that
    # tool_assign_heading_styles does NOT auto-promote any other
    # paragraph (e.g. the TOC entry) to Heading 1 during mode=full.
    doc.add_paragraph("引言", style="Heading 1")
    doc.add_paragraph("正文段落示例。", style="Normal")
    # toc.entry_count: 1 TOC entry to match the 1 heading.
    doc.add_paragraph("引言 ............ 1", style="TOC 1")

    doc.save(path)


def build_messy_docx(path: str) -> None:
    """A docx that intentionally violates ≥ 3 SCAU rules.

    Violations introduced (deliberately, for the sad-path test):
    - body.line_spacing: Normal style set to 2.0 (expected 1.5)
    - heading.h1.style_present: no paragraph uses Heading 1
    - toc.entry_count: 2 TOC entries with 0 headings → mismatch
    - page.margin.* (×4): default python-docx margins (≈2.54cm) ≠ 2.4cm
    - front_matter.*.present (×4): no abstract / keywords paragraphs

    The rules NOT violated (so the test can still distinguish "noise"
    from intentional violations) are body.font.east_asia and
    body.font.size — those stay at SCAU defaults so the report shows
    a clear pattern of failures, not a wall of red.
    """
    doc = Document()

    # Skip _apply_scau_page_layout: default python-docx margins (≈2.54cm
    # / 3.17cm) violate the SCAU 2.4cm spec. This is the page.* group
    # of violations.

    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    _set_east_asia_font(normal, "宋体")
    # ❌ body.line_spacing
    normal.paragraph_format.line_spacing = 2.0
    # body.first_line_indent intentionally LEFT unset → also fails

    _ensure_toc1_style(doc)

    # ❌ heading.h1.style_present: no Heading 1 anywhere
    doc.add_paragraph("绪论小标题（未套用样式）", style="Normal")
    doc.add_paragraph("正文段落。", style="Normal")

    # ❌ toc.entry_count: 2 TOC entries with 0 headings
    doc.add_paragraph("引言 ............ 1", style="TOC 1")
    doc.add_paragraph("结论 ............ 5", style="TOC 1")

    doc.save(path)


if __name__ == "__main__":  # pragma: no cover
    import sys

    out_dir = sys.argv[1] if len(sys.argv) > 1 else "."
    build_perfect_docx(f"{out_dir}/scau_perfect_thesis.docx")
    build_messy_docx(f"{out_dir}/scau_messy_thesis.docx")
    print("fixtures written under", out_dir)
