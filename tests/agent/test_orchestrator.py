"""Orchestrator — snapshot / policies / planner / harness / diagnoser / delivery.

Combined into one file because the dependencies form a single graph and
sharing fixtures keeps each test under 20 lines.

Backs requirements:
- R6.1 ~ R6.10  — main loop, exit conditions, modes, rollback, timeout
- R5.5 ~ R5.7   — diagnoser confidence handling, no-LLM fallback
- R7.1 ~ R7.5   — four artefacts, four-bucket status mapping
- R11.1 ~ R11.8 — trace shape and log level
- R13.1 / R13.2 — input not overwritten
"""

import json
import tempfile
import time
import unittest
from pathlib import Path

from docx import Document


# ---------------------------------------------------------------------------
# Snapshot
# ---------------------------------------------------------------------------

class SnapshotTests(unittest.TestCase):
    def test_take_creates_file_and_lru_evicts(self):
        from thesis_agent.ingest.document_model import DocumentModel
        from thesis_agent.orchestrator.snapshot import SnapshotManager

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "x.docx"
            d = Document()
            d.add_paragraph("hi")
            d.save(path)
            dm = DocumentModel.from_path(str(path))
            sm = SnapshotManager(work_dir=tmp, capacity=2)
            for i in range(5):
                sm.take(dm, tool_name=f"t{i}")
            files = list(Path(sm.directory).glob("*.docx"))
            self.assertEqual(len(files), 2)

    def test_rollback_last_restores_state(self):
        from thesis_agent.ingest.document_model import DocumentModel
        from thesis_agent.orchestrator.snapshot import SnapshotManager

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "x.docx"
            d = Document()
            d.add_paragraph("原文")
            d.save(path)
            dm = DocumentModel.from_path(str(path))
            sm = SnapshotManager(work_dir=tmp, capacity=2)
            sm.take(dm, tool_name="before_edit")
            with dm.write() as w:
                w.set_paragraph_text(0, "已改")
            self.assertEqual(dm.paragraphs()[0].text, "已改")
            self.assertTrue(sm.rollback_last(dm))
            self.assertEqual(dm.paragraphs()[0].text, "原文")


# ---------------------------------------------------------------------------
# Policies
# ---------------------------------------------------------------------------

class PoliciesTests(unittest.TestCase):
    def test_all_must_pass_true_when_only_should_fails(self):
        from thesis_agent.evaluators.types import CheckResult, EvalReport
        from thesis_agent.orchestrator.policies import all_must_pass

        report = EvalReport(profile="t", results=[
            CheckResult("a", "pass", "", {}, "must"),
            CheckResult("b", "fail", "", {}, "should"),
        ])
        self.assertTrue(all_must_pass(report))

    def test_plans_equivalent_uses_canonical_json(self):
        from thesis_agent.orchestrator.planner import Step
        from thesis_agent.orchestrator.policies import plans_equivalent

        a = [Step(tool="x", params={"a": 1, "b": 2})]
        b = [Step(tool="x", params={"b": 2, "a": 1})]
        self.assertTrue(plans_equivalent(a, b))

    def test_should_exit_when_max_iterations_reached(self):
        from thesis_agent.evaluators.types import CheckResult, EvalReport
        from thesis_agent.orchestrator.policies import (
            Policies, should_exit,
        )

        # Need at least one failing must-severity rule so the
        # "all must pass" exit doesn't fire first.
        rep = EvalReport(profile="t", results=[
            CheckResult("a", "fail", "", {}, "must"),
        ])
        decision = should_exit(
            eval_report=rep, iteration=3,
            prev_plan=None, new_plan=[],
            deadline=time.monotonic() + 100, cancelled=False,
            policies=Policies(max_iterations=3),
        )
        self.assertTrue(decision.should_exit)
        self.assertIn("max_iterations", decision.reason)


# ---------------------------------------------------------------------------
# Planner
# ---------------------------------------------------------------------------

class PlannerTests(unittest.TestCase):
    def test_default_plan_full_mode_starts_with_assign_then_body(self):
        from thesis_agent.orchestrator.planner import default_plan

        plan = default_plan(rule_set=None, mode="full")
        self.assertGreater(len(plan), 0)
        self.assertEqual(plan[0].tool, "tool_assign_heading_styles")
        self.assertEqual(plan[1].tool, "tool_format_body")
        self.assertIn("tool_normalize_heading_spacing", [s.tool for s in plan])

    def test_replan_appends_diagnosis_fix_plan(self):
        from thesis_agent.diagnoser.types import Diagnosis, ToolCall
        from thesis_agent.orchestrator.planner import replan

        d = Diagnosis(
            rule_id="body.line_spacing",
            root_cause="...",
            fix_plan=[ToolCall(tool="tool_format_body", params={"line_spacing": 1.5})],
            confidence=0.9,
            needs_human=False,
            rationale="",
        )
        new_plan = replan([d], prev_plan=[])
        self.assertEqual([s.tool for s in new_plan], ["tool_format_body"])

    def test_replan_returns_equivalent_when_all_human(self):
        from thesis_agent.diagnoser.types import Diagnosis
        from thesis_agent.orchestrator.planner import replan

        prev = []
        d = Diagnosis(
            rule_id="x", root_cause="", fix_plan=[],
            confidence=0.0, needs_human=True, rationale="LLM 缺失",
        )
        self.assertEqual(replan([d], prev), prev)

    def test_replan_falls_back_to_rule_fix_tool_when_llm_has_no_plan(self):
        from thesis_agent.diagnoser.types import Diagnosis
        from thesis_agent.evaluators.types import CheckResult, EvalReport
        from thesis_agent.orchestrator.planner import replan
        from thesis_agent.spec.rule_set import Rule, RuleSet

        rule_set = RuleSet(
            profile="t",
            version="1",
            rules=[
                Rule(
                    id="body.line_spacing",
                    scope="style",
                    locator={"style_name": "Normal"},
                    predicate="equals",
                    expected=1.5,
                    severity="must",
                    fix_tool="tool_format_body",
                    fix_params_template={"line_spacing": "{expected}"},
                ),
            ],
        )
        eval_report = EvalReport(profile="t", results=[
            CheckResult(
                "body.line_spacing", "fail", "actual=2.0 expected=1.5",
                {"style_name": "Normal"}, "must",
            ),
        ])
        diagnosis = Diagnosis(
            rule_id="body.line_spacing",
            root_cause="",
            fix_plan=[],
            confidence=0.0,
            needs_human=True,
            rationale="未配置 LLM",
        )

        new_plan = replan(
            [diagnosis],
            prev_plan=[],
            eval_report=eval_report,
            rule_set=rule_set,
            allow_rule_fallback=True,
        )

        self.assertEqual(len(new_plan), 1)
        self.assertEqual(new_plan[0].tool, "tool_format_body")
        self.assertEqual(new_plan[0].params, {"line_spacing": 1.5})


# ---------------------------------------------------------------------------
# Diagnoser
# ---------------------------------------------------------------------------

class DiagnoserTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches

        reset_caches()

    def test_no_llm_returns_needs_human(self):
        from thesis_agent.diagnoser.diagnoser import diagnose
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("a", "fail", "ev", {}, "must"),
        ])
        out = diagnose(rep, doc=None, llm=None)
        self.assertEqual(len(out), 1)
        self.assertTrue(out[0].needs_human)

    def test_mock_llm_canned_response_used(self):
        from thesis_agent.diagnoser.diagnoser import diagnose
        from thesis_agent.diagnoser.llm_client import MockLLMClient
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        llm = MockLLMClient(canned={
            "body.line_spacing": {
                "rule_id": "body.line_spacing",
                "root_cause": "fixed 20pt",
                "fix_plan": [{"tool": "tool_format_body",
                              "params": {"line_spacing": 1.5}}],
                "confidence": 0.9,
                "needs_human": False,
                "rationale": "",
            }
        })
        rep = EvalReport(profile="t", results=[
            CheckResult("body.line_spacing", "fail", "ev", {}, "must"),
        ])
        out = diagnose(rep, doc=None, llm=llm)
        self.assertEqual(out[0].fix_plan[0].tool, "tool_format_body")

    def test_repeat_fail_downgrades_confidence(self):
        from thesis_agent.diagnoser.diagnoser import diagnose, reset_caches
        from thesis_agent.diagnoser.llm_client import MockLLMClient
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        reset_caches()
        llm = MockLLMClient(canned={
            "body.line_spacing": {
                "rule_id": "body.line_spacing",
                "root_cause": "x",
                "fix_plan": [],
                "confidence": 0.95,
                "needs_human": False,
                "rationale": "",
            }
        })
        cr = CheckResult("body.line_spacing", "fail", "ev", {}, "must")
        # First diagnose: confidence preserved (cached straight from llm)
        rep = EvalReport(profile="t", results=[cr])
        first = diagnose(rep, doc=None, llm=llm)
        # Second invocation hits cache, so simulate a repeat by clearing
        # only the response cache but keeping repeat-fail history.
        from thesis_agent.diagnoser import diagnoser as mod
        mod._DIAGNOSIS_CACHE.clear()
        second = diagnose(EvalReport(profile="t", results=[cr]), doc=None, llm=llm)
        self.assertLessEqual(second[0].confidence, 0.5)
        self.assertTrue(second[0].needs_human)


# ---------------------------------------------------------------------------
# Trace
# ---------------------------------------------------------------------------

class TraceTests(unittest.TestCase):
    def test_records_one_json_per_line(self):
        from thesis_agent.delivery.trace import Trace

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "trace.jsonl"
            tr = Trace(str(path))
            tr.record(kind="plan", payload={"x": 1})
            tr.record(kind="eval", payload={"y": 2})
            lines = path.read_text(encoding="utf-8").splitlines()
            self.assertEqual(len(lines), 2)
            for line in lines:
                obj = json.loads(line)
                self.assertIn("ts", obj)
                self.assertIn("kind", obj)
                self.assertIn("payload", obj)

    def test_invalid_kind_rejected(self):
        from thesis_agent.delivery.trace import (
            InvalidTraceKindError, Trace,
        )

        with tempfile.TemporaryDirectory() as tmp:
            tr = Trace(str(Path(tmp) / "trace.jsonl"))
            with self.assertRaises(InvalidTraceKindError):
                tr.record(kind="random", payload={})

    def test_llm_payloads_dropped_at_info_level(self):
        from thesis_agent.delivery.trace import Trace

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "trace.jsonl"
            tr = Trace(str(path), log_level="INFO")
            tr.record(kind="llm_request", payload={"prompt": "x"})
            self.assertEqual(path.read_text(encoding="utf-8"), "")


# ---------------------------------------------------------------------------
# Report — status mapping (R7.4)
# ---------------------------------------------------------------------------

class ReportStatusMappingTests(unittest.TestCase):
    def test_pass_maps_to_done(self):
        from thesis_agent.delivery.report import build_delivery_report
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("a", "pass", "", {}, "must"),
        ])
        delivery = build_delivery_report(
            rule_set=None, eval_report=rep, diagnoses=[],
            mode="full", iterations=1, exit_reason="ok",
        )
        self.assertEqual(delivery.summary, {"total": 1, "done": 1,
                                            "partial": 0, "failed": 0,
                                            "skipped": 0})

    def test_fail_with_human_diagnosis_maps_to_failed(self):
        from thesis_agent.delivery.report import build_delivery_report
        from thesis_agent.diagnoser.types import Diagnosis
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("a", "fail", "ev", {}, "must"),
        ])
        d = Diagnosis(rule_id="a", root_cause="", fix_plan=[],
                      confidence=0.0, needs_human=True, rationale="")
        delivery = build_delivery_report(
            rule_set=None, eval_report=rep, diagnoses=[d],
            mode="full", iterations=1, exit_reason="ok",
        )
        self.assertEqual(delivery.summary["failed"], 1)
        self.assertEqual(delivery.summary["partial"], 0)

    def test_error_maps_to_failed(self):
        from thesis_agent.delivery.report import build_delivery_report
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("a", "error", "locator gone", {}, "must"),
        ])
        delivery = build_delivery_report(
            rule_set=None, eval_report=rep, diagnoses=[],
            mode="full", iterations=1, exit_reason="ok",
        )
        self.assertEqual(delivery.summary["failed"], 1)

    def test_fail_with_tool_attempt_maps_to_partial(self):
        from thesis_agent.delivery.report import build_delivery_report
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("body.line_spacing", "fail", "still wrong", {}, "must"),
        ])
        delivery = build_delivery_report(
            rule_set=None,
            eval_report=rep,
            diagnoses=[],
            mode="full",
            iterations=2,
            exit_reason="converged",
            fix_attempts_by_rule={
                "body.line_spacing": [
                    {
                        "iteration": 0,
                        "tool": "tool_format_body",
                        "params": {"line_spacing": 1.5},
                        "ok": True,
                        "message": "body style updated",
                    }
                ]
            },
        )

        self.assertEqual(delivery.items[0].status, "partial")
        self.assertEqual(delivery.summary["partial"], 1)
        self.assertEqual(delivery.summary["failed"], 0)

    def test_build_delivery_report_carries_fix_attempts_by_rule(self):
        from thesis_agent.delivery.report import build_delivery_report
        from thesis_agent.evaluators.types import CheckResult, EvalReport

        rep = EvalReport(profile="t", results=[
            CheckResult("body.line_spacing", "pass", "actual=1.5", {}, "must"),
        ])
        delivery = build_delivery_report(
            rule_set=None,
            eval_report=rep,
            diagnoses=[],
            mode="full",
            iterations=2,
            exit_reason="ok",
            fix_attempts_by_rule={
                "body.line_spacing": [
                    {
                        "iteration": 0,
                        "tool": "tool_format_body",
                        "params": {"line_spacing": 1.5},
                        "ok": True,
                        "message": "body style updated",
                    }
                ]
            },
        )

        item = delivery.items[0]
        self.assertEqual(item.fix_attempts[0]["tool"], "tool_format_body")
        self.assertTrue(item.fix_attempts[0]["ok"])


# ---------------------------------------------------------------------------
# Harness — happy path
# ---------------------------------------------------------------------------

def _make_simple_docx(path, paras):
    doc = Document()
    for text, style in paras:
        p = doc.add_paragraph(text)
        if style:
            p.style = doc.styles[style]
    doc.save(path)


class HarnessTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_eval_only_does_not_modify_input(self):
        from thesis_agent.orchestrator.harness import RunOptions, run

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            _make_simple_docx(input_path, [("第一章", "Heading 1"), ("正文", None)])
            before = input_path.read_bytes()
            result = run(
                input_path=str(input_path),
                profile="scau_2024",
                mode="eval_only",
                options=RunOptions(output_dir=tmp),
            )
            self.assertTrue(result.ok)
            self.assertIsNone(result.docx_path)
            self.assertEqual(input_path.read_bytes(), before)
            # Three artefacts present
            self.assertTrue(Path(result.report_md_path).exists())
            self.assertTrue(Path(result.report_json_path).exists())
            self.assertTrue(Path(result.trace_path).exists())

    def test_overwrite_input_path_raises(self):
        from thesis_agent.orchestrator.harness import (
            OverwriteInputError, RunOptions, run,
        )

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "x.docx"
            _make_simple_docx(input_path, [("hi", None)])
            with self.assertRaises(OverwriteInputError):
                run(
                    input_path=str(input_path),
                    profile="scau_2024",
                    mode="full",
                    options=RunOptions(output_path=str(input_path)),
                )

    def test_invalid_mode_rejected(self):
        from thesis_agent.orchestrator.harness import InvalidModeError, run

        with self.assertRaises(InvalidModeError):
            run(
                input_path="any",
                profile="scau_2024",
                mode="bogus",  # type: ignore[arg-type]
            )


if __name__ == "__main__":
    unittest.main()
