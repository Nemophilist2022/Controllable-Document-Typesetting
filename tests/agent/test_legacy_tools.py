"""Contract + smoke tests for the 12 legacy-wrapping Tools added in A1.

The 12 Tools are thin wrappers around existing ``thesis_formatter/*``
functions, which themselves are exercised by the project's pre-existing
test suite. This file only verifies:

1. Each Tool satisfies the Tool protocol (R3.2).
2. Auto-discovery picks them up through the registry.
3. ``ToolContext.config`` carries the deep-merged profile dict (C3 fix).
4. Each Tool returns a ``ToolResult(ok=False, ...)`` rather than raising
   when its underlying function blows up.
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from thesis_agent.ingest.document_model import DocumentModel
from thesis_agent.spec.profiles import load_profile
from thesis_agent.tools import registry
from thesis_agent.tools.base import ToolContext, is_tool


_LEGACY_TOOL_NAMES = {
    "tool_renumber_headings",
    "tool_normalize_heading_spacing",
    "tool_setup_multilevel_list",
    "tool_format_figure_captions",
    "tool_format_table_captions",
    "tool_format_three_line_tables",
    "tool_format_references",
    "tool_setup_headers",
    "tool_normalize_sections",
    "tool_setup_page_numbers",
    "tool_setup_page_numbers_strict",
    "tool_insert_cover_and_declaration",
}


class _FakeSnapshotMgr:
    def take(self, doc, tool_name=""):
        return f"snap-{tool_name}"


def _make_ctx(config):
    return ToolContext(
        trace=None,
        snapshot_mgr=_FakeSnapshotMgr(),
        config=config or {},
        runtime={},
    )


def _make_simple_dm(tmp: Path) -> DocumentModel:
    path = tmp / "simple.docx"
    doc = Document()
    doc.add_paragraph("第一章 绪论", style="Heading 1")
    doc.add_paragraph("正文")
    doc.save(path)
    return DocumentModel.from_path(str(path))


class LegacyToolContractTests(unittest.TestCase):
    def setUp(self):
        registry.clear()
        registry.autoload()

    def test_all_legacy_tools_registered(self):
        names = {t.name for t in registry.all_tools()}
        for expected in _LEGACY_TOOL_NAMES:
            with self.subTest(tool=expected):
                self.assertIn(expected, names)

    def test_all_legacy_tools_satisfy_protocol(self):
        for name in _LEGACY_TOOL_NAMES:
            with self.subTest(tool=name):
                tool = registry.get(name)
                self.assertTrue(is_tool(tool))
                self.assertIsInstance(tool.input_schema, dict)
                self.assertIsInstance(tool.requires, list)
                self.assertIsInstance(tool.idempotent, bool)


class HarnessConfigPlumbingTests(unittest.TestCase):
    """C3 fix: harness must pipe the profile's source config into
    ``ToolContext.config`` so legacy wrappers see the same cfg shape
    they expect."""

    def test_rule_set_metadata_carries_source_config(self):
        rs = load_profile("scau_2024")
        cfg = rs.metadata.get("source_config")
        self.assertIsInstance(cfg, dict)
        # Sanity: it must look like a real profile config.
        self.assertIn("body", cfg)
        self.assertIn("fonts", cfg)
        self.assertIn("page", cfg)


class LegacyToolSmokeTests(unittest.TestCase):
    """Each Tool is a pass-through to a legacy function. We don't
    re-verify the legacy semantics here (they have their own tests);
    we only confirm the wrapper exits cleanly with ok=True or ok=False
    instead of letting an exception propagate."""

    def setUp(self):
        registry.clear()
        registry.autoload()
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.dm = _make_simple_dm(self.tmp)
        self.cfg = load_profile("scau_2024").metadata["source_config"]
        self.ctx = _make_ctx(self.cfg)

    def tearDown(self):
        self._tmp.cleanup()

    def _expect_no_raise(self, tool_name: str):
        tool = registry.get(tool_name)
        result = tool.run(self.dm, {}, self.ctx)
        # A wrapped Tool must always answer with a ToolResult; the ok
        # flag may be either way (some legacy fns require richer fixtures
        # to succeed) but the call itself must not raise.
        self.assertIn(result.ok, (True, False), msg=f"{tool_name} returned {result!r}")
        # rollback_token comes from the snapshot manager and must be set
        # — proving the wrapper went through the snapshot path.
        self.assertIsNotNone(result.rollback_token, msg=f"{tool_name} skipped snapshot")

    def test_normalize_sections_does_not_raise(self):
        self._expect_no_raise("tool_normalize_sections")

    def test_setup_headers_does_not_raise(self):
        self._expect_no_raise("tool_setup_headers")

    def test_format_three_line_tables_does_not_raise(self):
        self._expect_no_raise("tool_format_three_line_tables")

    def test_format_figure_captions_does_not_raise(self):
        self._expect_no_raise("tool_format_figure_captions")

    def test_format_table_captions_does_not_raise(self):
        self._expect_no_raise("tool_format_table_captions")

    def test_format_references_does_not_raise(self):
        self._expect_no_raise("tool_format_references")

    def test_renumber_headings_does_not_raise(self):
        self._expect_no_raise("tool_renumber_headings")

    def test_normalize_heading_spacing_does_not_raise(self):
        self._expect_no_raise("tool_normalize_heading_spacing")

    def test_setup_multilevel_list_does_not_raise(self):
        self._expect_no_raise("tool_setup_multilevel_list")

    def test_setup_page_numbers_strict_does_not_raise(self):
        self._expect_no_raise("tool_setup_page_numbers_strict")


class LegacyToolFailureTests(unittest.TestCase):
    """Failure path: a Tool whose wrapped function raises must surface
    that as ok=False rather than propagating."""

    def setUp(self):
        registry.clear()
        registry.autoload()

    def test_setup_page_numbers_with_empty_config_returns_ok_false(self):
        """``setup_page_numbers`` reads ``cfg['page_numbers']``; an empty
        cfg should make it raise, which the wrapper must catch."""
        with tempfile.TemporaryDirectory() as tmp:
            dm = _make_simple_dm(Path(tmp))
            tool = registry.get("tool_setup_page_numbers")
            ctx = _make_ctx({})  # intentionally empty
            result = tool.run(dm, {}, ctx)
            self.assertFalse(result.ok)
            self.assertIn("tool_setup_page_numbers failed", result.message)


if __name__ == "__main__":
    unittest.main()
