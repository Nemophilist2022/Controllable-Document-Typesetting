"""v0.2 — implant check coverage for previously-skipped rules.

Targets:
- ``heading.h{1..4}.font.east_asia`` / ``font.size`` 通过推断 attr 走 check_styles
- ``caption.font.east_asia`` / ``font.size``        — check_paragraphs
- ``reference.first_line_indent``                   — check_paragraphs
- ``header.enabled``                                — check_doc
- ``toc.font.east_asia``                            — check_paragraphs

These tests bypass the harness — we just construct a docx, build the
rule directly, and call the relevant check function.
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from thesis_agent.evaluators.checks.check_doc import check_header_enabled
from thesis_agent.evaluators.checks.check_paragraphs import (
    check_paragraph_group_attr,
)
from thesis_agent.evaluators.checks.check_styles import (
    _infer_attr,
    check_style_attr,
)
from thesis_agent.spec.rule_set import Rule


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_east_asia_on_style(style, name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _set_east_asia_on_run(run, name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _make_doc_with_heading():
    doc = Document()
    style = doc.styles["Heading 1"]
    style.font.size = Pt(14)
    _set_east_asia_on_style(style, "黑体")
    doc.add_paragraph("第一章", style="Heading 1")
    doc.add_paragraph("正文", style="Normal")
    return doc


def _heading_rule(rule_id: str, *, expected, predicate="equals", style="Heading 1"):
    return Rule(
        id=rule_id,
        scope="style",
        locator={"style_name": style},  # 注意：deliberately 不带 attr
        predicate=predicate,
        expected=expected,
        severity="should",
    )


# ---------------------------------------------------------------------------
# check_styles — heading rules now resolve attr by rule.id suffix
# ---------------------------------------------------------------------------

class HeadingStyleAttrInferenceTests(unittest.TestCase):
    def test_infer_attr_from_id_east_asia(self):
        self.assertEqual(_infer_attr("heading.h1.font.east_asia"),
                         "east_asia_font")

    def test_infer_attr_from_id_size(self):
        self.assertEqual(_infer_attr("heading.h2.font.size"), "size_pt")

    def test_infer_attr_from_id_bold(self):
        self.assertEqual(_infer_attr("heading.h3.bold"), "bold")

    def test_unknown_suffix_returns_none(self):
        self.assertIsNone(_infer_attr("heading.h1.something_strange"))


class HeadingFontPassFailTests(unittest.TestCase):
    def test_h1_east_asia_passes_when_style_matches(self):
        doc = _make_doc_with_heading()
        rule = _heading_rule(
            "heading.h1.font.east_asia", expected="黑体"
        )
        result = check_style_attr(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)

    def test_h1_east_asia_fails_when_style_mismatches(self):
        doc = _make_doc_with_heading()
        rule = _heading_rule(
            "heading.h1.font.east_asia", expected="宋体"
        )
        result = check_style_attr(rule, doc)
        self.assertEqual(result.status, "fail", msg=result.evidence)
        self.assertIn("黑体", result.evidence)

    def test_h1_size_passes(self):
        doc = _make_doc_with_heading()
        rule = _heading_rule("heading.h1.font.size", expected=14)
        result = check_style_attr(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)


# ---------------------------------------------------------------------------
# check_paragraphs — caption / reference / toc_entries
# ---------------------------------------------------------------------------

def _para_rule(rule_id, locator, *, expected):
    return Rule(
        id=rule_id,
        scope="paragraph",
        locator=locator,
        predicate="equals",
        expected=expected,
        severity="should",
    )


class CaptionFontTests(unittest.TestCase):
    def test_caption_east_asia_pass(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("图 1 示意图")
        _set_east_asia_on_run(run, "宋体")
        rule = _para_rule(
            "caption.font.east_asia",
            {"caption": True, "attr": "east_asia_font"},
            expected="宋体",
        )
        result = check_paragraph_group_attr(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)

    def test_caption_east_asia_fail(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("图 1 示意图")
        _set_east_asia_on_run(run, "黑体")
        rule = _para_rule(
            "caption.font.east_asia",
            {"caption": True, "attr": "east_asia_font"},
            expected="宋体",
        )
        result = check_paragraph_group_attr(rule, doc)
        self.assertEqual(result.status, "fail", msg=result.evidence)

    def test_caption_skip_when_no_captions(self):
        doc = Document()
        doc.add_paragraph("普通正文")
        rule = _para_rule(
            "caption.font.east_asia",
            {"caption": True, "attr": "east_asia_font"},
            expected="宋体",
        )
        result = check_paragraph_group_attr(rule, doc)
        self.assertEqual(result.status, "skip")


class ReferenceIndentTests(unittest.TestCase):
    def test_reference_first_line_indent_pass(self):
        doc = Document()
        para = doc.add_paragraph("[1] Smith J. xxx")
        para.paragraph_format.first_line_indent = Pt(-24)
        rule = _para_rule(
            "reference.first_line_indent",
            {"references_section": True, "attr": "first_line_indent_pt"},
            expected=-24,
        )
        result = check_paragraph_group_attr(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)


class TocFontTests(unittest.TestCase):
    def test_toc_east_asia_pass(self):
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()
        if "TOC 1" not in [s.name for s in doc.styles]:
            doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
        para = doc.add_paragraph("引言 ............ 1", style="TOC 1")
        run = para.add_run("")  # ensure at least one run
        _set_east_asia_on_run(para.runs[0], "宋体")
        rule = _para_rule(
            "toc.font.east_asia",
            {"toc_entries": True, "attr": "east_asia_font"},
            expected="宋体",
        )
        result = check_paragraph_group_attr(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)


# ---------------------------------------------------------------------------
# check_doc — header.enabled
# ---------------------------------------------------------------------------

class HeaderEnabledTests(unittest.TestCase):
    def test_header_disabled_passes_when_no_header_text(self):
        doc = Document()
        rule = Rule(
            id="header.enabled",
            scope="doc",
            locator={"header": True, "attr": "enabled"},
            predicate="equals",
            expected=False,
            severity="info",
        )
        result = check_header_enabled(rule, doc)
        self.assertEqual(result.status, "pass", msg=result.evidence)

    def test_header_enabled_fails_when_no_header_but_expected(self):
        doc = Document()
        rule = Rule(
            id="header.enabled",
            scope="doc",
            locator={"header": True, "attr": "enabled"},
            predicate="equals",
            expected=True,
            severity="info",
        )
        result = check_header_enabled(rule, doc)
        self.assertEqual(result.status, "fail")

    def test_header_enabled_passes_when_header_present(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "PAPER"
        rule = Rule(
            id="header.enabled",
            scope="doc",
            locator={"header": True, "attr": "enabled"},
            predicate="equals",
            expected=True,
            severity="info",
        )
        result = check_header_enabled(rule, doc)
        self.assertEqual(result.status, "pass")


if __name__ == "__main__":
    unittest.main()
