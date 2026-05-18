"""Wiring tests: harness builds the right LLM client based on options
and surfaces telemetry through report.json.meta.

NOTE on environment patching: we never use ``clear=True`` on
``mock.patch.dict("os.environ", ...)`` because the project's win32com
dependency reads system env vars (``WINDIR`` etc.) at import time and
crashes when they vanish. Instead we explicitly remove only the
``THESIS_AGENT_LLM_*`` keys.
"""

import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document


def _make_doc(path):
    d = Document()
    d.add_paragraph("第一章", style="Heading 1")
    d.add_paragraph("正文")
    d.save(path)


def _no_llm_env() -> dict[str, str]:
    """Return a copy of os.environ with all LLM-related keys removed."""
    env = dict(os.environ)
    for k in list(env):
        if k.startswith("THESIS_AGENT_LLM_"):
            env.pop(k, None)
    return env


def _llm_env(**overrides) -> dict[str, str]:
    env = _no_llm_env()
    env.update(overrides)
    return env


class HarnessLLMWiringTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_no_credentials_means_no_llm(self):
        from thesis_agent.orchestrator.harness import RunOptions, _build_llm_client

        with mock.patch.dict("os.environ", _no_llm_env(), clear=True):
            self.assertIsNone(_build_llm_client(RunOptions()))

    def test_explicit_credentials_build_client(self):
        from thesis_agent.diagnoser.openai_client import OpenAICompatibleClient
        from thesis_agent.orchestrator.harness import RunOptions, _build_llm_client

        with mock.patch.dict("os.environ", _no_llm_env(), clear=True):
            client = _build_llm_client(RunOptions(
                llm_api_key="k",
                llm_model="gpt-4o-mini",
                llm_base_url="https://api.example.com/v1",
            ))
        self.assertIsInstance(client, OpenAICompatibleClient)

    def test_no_llm_flag_overrides_env(self):
        from thesis_agent.orchestrator.harness import RunOptions, _build_llm_client

        env = _llm_env(THESIS_AGENT_LLM_API_KEY="k")
        with mock.patch.dict("os.environ", env, clear=True):
            self.assertIsNone(_build_llm_client(RunOptions(llm_disabled=True)))


class ReportLLMTelemetryTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_eval_only_run_without_llm_meta_has_zero_calls(self):
        from thesis_agent.orchestrator.harness import RunOptions, run

        with tempfile.TemporaryDirectory() as tmp:
            in_path = Path(tmp) / "x.docx"
            _make_doc(in_path)
            with mock.patch.dict("os.environ", _no_llm_env(), clear=True):
                result = run(
                    input_path=str(in_path),
                    profile="scau_2024",
                    mode="eval_only",
                    options=RunOptions(output_dir=tmp),
                )
            report = json.loads(Path(result.report_json_path).read_text(encoding="utf-8"))
            meta = report["meta"]
            self.assertEqual(meta["llm_calls_count"], 0)
            self.assertEqual(meta["llm_cost_estimate_usd"], 0.0)
            self.assertEqual(meta["llm_telemetry"], {})


class CLILLMFlagsTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_no_llm_flag_disables_client(self):
        """``--no-llm`` must disable the client even if env has a key."""
        import io
        from contextlib import redirect_stdout

        from thesis_agent.cli import main

        env = _llm_env(THESIS_AGENT_LLM_API_KEY="would-be-used")
        with mock.patch.dict("os.environ", env, clear=True):
            with tempfile.TemporaryDirectory() as tmp:
                in_path = Path(tmp) / "x.docx"
                _make_doc(in_path)
                buf = io.StringIO()
                with redirect_stdout(buf):
                    rc = main([
                        "run",
                        "--input", str(in_path),
                        "--profile", "scau_2024",
                        "--mode", "eval_only",
                        "--output-dir", tmp,
                        "--no-llm",
                    ])
                self.assertEqual(rc, 0)
                report = json.loads(
                    (Path(tmp) / "x_report.json").read_text(encoding="utf-8")
                )
                self.assertEqual(report["meta"]["llm_calls_count"], 0)


if __name__ == "__main__":
    unittest.main()
