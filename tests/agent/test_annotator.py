"""Annotator — paragraph shading for partial / failed delivery items."""

import tempfile
import unittest
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from thesis_agent.delivery.annotator import (
    ANNOTATION_FILL, AnnotationResult, annotate,
)


# ---------------------------------------------------------------------------
# Lightweight delivery stand-in
# ---------------------------------------------------------------------------

@dataclass
class _Item:
    rule_id: str
    status: str
    locator: dict
    severity: str = "must"
    evidence: str = "ev"


@dataclass
class _Delivery:
    items: list


def _make_doc(path: Path, paragraphs):
    doc = Document()
    for text, style in paragraphs:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    doc.save(path)


def _shading_fill(para) -> str | None:
    p_pr = para._element.find(qn("w:pPr"))
    if p_pr is None:
        return None
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class AnnotateByParagraphIndexTests(unittest.TestCase):
    def test_failed_item_with_paragraph_index_gets_shaded(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [
                ("第一段", None),
                ("第二段需要标注", None),
                ("第三段", None),
            ])
            delivery = _Delivery(items=[
                _Item("body.line_spacing", "failed", {"paragraph_index": 1}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 1)
            self.assertFalse(result.warnings, msg=result.warnings)
            doc = Document(str(p))
            self.assertEqual(_shading_fill(doc.paragraphs[1]), ANNOTATION_FILL)
            # First paragraph must remain unmarked.
            self.assertIsNone(_shading_fill(doc.paragraphs[0]))
            # Inline note appended after original text.
            self.assertIn("body.line_spacing", doc.paragraphs[1].text)


class AnnotateByStyleNameTests(unittest.TestCase):
    def test_heading_style_locator_shades_every_heading_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [
                ("引言", "Heading 1"),
                ("正文", None),
                ("参考文献", "Heading 1"),
            ])
            delivery = _Delivery(items=[
                _Item("heading.h1.font.east_asia", "failed",
                      {"style_name": "Heading 1", "attr": "east_asia_font"}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 2)
            doc = Document(str(p))
            self.assertEqual(_shading_fill(doc.paragraphs[0]), ANNOTATION_FILL)
            self.assertIsNone(_shading_fill(doc.paragraphs[1]))
            self.assertEqual(_shading_fill(doc.paragraphs[2]), ANNOTATION_FILL)

    def test_normal_style_locator_is_NOT_marked_doc_wide(self):
        """Marking every Normal paragraph would be visual noise. The
        annotator deliberately skips locator['style_name']=='Normal'."""
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [
                ("段一", None),  # Normal
                ("段二", None),
            ])
            delivery = _Delivery(items=[
                _Item("body.line_spacing", "failed", {"style_name": "Normal"}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 0)


class AnnotateFrontMatterTests(unittest.TestCase):
    def test_cn_keywords_locator_finds_paragraph_by_prefix(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [
                ("摘要", None),
                ("正文略", None),
                ("关键词：A; B; C", None),
            ])
            delivery = _Delivery(items=[
                _Item("front_matter.cn_keywords.present", "failed",
                      {"front_matter": "cn_keywords"}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 1)
            doc = Document(str(p))
            self.assertEqual(_shading_fill(doc.paragraphs[2]), ANNOTATION_FILL)


class AnnotateMergesPerParagraphTests(unittest.TestCase):
    def test_two_rules_on_same_paragraph_share_one_shading(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [
                ("引言", "Heading 1"),
            ])
            delivery = _Delivery(items=[
                _Item("rule.a", "failed", {"paragraph_index": 0}),
                _Item("rule.b", "partial", {"paragraph_index": 0}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 1)
            doc = Document(str(p))
            text = doc.paragraphs[0].text
            self.assertIn("rule.a", text)
            self.assertIn("rule.b", text)


class AnnotateSkipsUnpinnableLocatorsTests(unittest.TestCase):
    def test_all_sections_locator_is_skipped(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [("仅一段", None)])
            delivery = _Delivery(items=[
                _Item("page.margin.top", "failed",
                      {"all_sections": True, "attr": "top_margin_cm"}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 0)
            self.assertEqual(result.skipped_items, 1)


class AnnotateSkipsNonViolatingItemsTests(unittest.TestCase):
    def test_done_and_skipped_items_do_not_get_shaded(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "x.docx"
            _make_doc(p, [("段一", None)])
            delivery = _Delivery(items=[
                _Item("rule.a", "done", {"paragraph_index": 0}),
                _Item("rule.b", "skipped", {"paragraph_index": 0}),
            ])
            result = annotate(str(p), delivery)
            self.assertEqual(result.annotated_paragraphs, 0)
            doc = Document(str(p))
            self.assertIsNone(_shading_fill(doc.paragraphs[0]))


class AnnotateFailureModeTests(unittest.TestCase):
    def test_missing_file_returns_warning_not_raise(self):
        result = annotate("/no/such/path.docx", _Delivery(items=[
            _Item("rule.x", "failed", {"paragraph_index": 0}),
        ]))
        self.assertIsInstance(result, AnnotationResult)
        self.assertEqual(result.annotated_paragraphs, 0)
        self.assertTrue(result.warnings)


# Note: a previous version of this file also exercised the full harness
# loop end-to-end. That test was removed because it pulled the Word COM
# import chain in (via thesis_runner → word_postprocess → win32com),
# which can deadlock when a stale winword.exe / gen_py cache is around.
# The 8 unit tests above already cover every annotator branch the
# harness depends on; the harness wiring itself is exercised by
# tests/agent/test_hitl_harness.py.


if __name__ == "__main__":
    unittest.main()
