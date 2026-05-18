"""DocumentModel — read API + controlled write API (R2.3, R2.4)."""

import tempfile
import unittest
from pathlib import Path

from docx import Document


def _make_doc(path, paragraphs=("段落一", "段落二", "段落三")):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


class DocumentModelReadAPITests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.path = self.tmp / "input.docx"
        _make_doc(self.path)

    def tearDown(self):
        self._tmp.cleanup()

    def test_from_path_loads_paragraphs(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        paras = dm.paragraphs()
        self.assertEqual(len(paras), 3)
        self.assertEqual(paras[0].text, "段落一")

    def test_paragraphs_returns_immutable_tuple(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        paras = dm.paragraphs()
        # tuple is immutable
        with self.assertRaises(TypeError):
            paras[0] = None  # type: ignore[index]

    def test_styles_view_can_iterate(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        names = {s.name for s in dm.styles()}
        self.assertIn("Normal", names)


class DocumentModelWriteAPITests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.path = self.tmp / "input.docx"
        _make_doc(self.path)

    def tearDown(self):
        self._tmp.cleanup()

    def test_set_paragraph_text_tracks_change(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        with dm.write() as w:
            w.set_paragraph_text(0, "新文本")
        self.assertIn(0, dm.last_changes.paragraphs)
        # Read API reflects the change
        self.assertEqual(dm.paragraphs()[0].text, "新文本")

    def test_save_writes_to_path(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        with dm.write() as w:
            w.set_paragraph_text(0, "已修改")
        out = self.tmp / "out.docx"
        dm.save(str(out))
        self.assertTrue(out.exists())
        # Round-trip via python-docx confirms persistence
        doc2 = Document(str(out))
        self.assertEqual(doc2.paragraphs[0].text, "已修改")

    def test_last_changes_resets_per_write_session(self):
        from thesis_agent.ingest.document_model import DocumentModel

        dm = DocumentModel.from_path(str(self.path))
        with dm.write() as w:
            w.set_paragraph_text(0, "A")
        self.assertEqual(dm.last_changes.paragraphs, [0])
        with dm.write() as w:
            w.set_paragraph_text(2, "B")
        # Second session should report only that session's changes
        self.assertEqual(dm.last_changes.paragraphs, [2])


if __name__ == "__main__":
    unittest.main()
