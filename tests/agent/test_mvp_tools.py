"""Contract + behaviour tests for the 4 MVP Tools (R3.2, R3.5, R3.7).

We bundle these into one file so the docx fixtures and fake snapshot
manager are written only once.
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document

from thesis_agent.ingest.document_model import DocumentModel
from thesis_agent.tools.base import ToolContext, ToolResult, is_tool


class _FakeSnapshotMgr:
    def __init__(self):
        self.take_calls = 0

    def take(self, doc, tool_name=""):
        self.take_calls += 1
        return f"snap-{self.take_calls}"


def _make_ctx(config=None):
    return ToolContext(
        trace=None,
        snapshot_mgr=_FakeSnapshotMgr(),
        config=config or {},
        runtime={},
    )


def _minimal_cfg():
    """A trimmed config sufficient for body / heading / toc tools."""
    from thesis_config import DEFAULT_CONFIG
    import copy

    return copy.deepcopy(DEFAULT_CONFIG)


def _make_docx_with_paragraphs(path, paragraphs):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


# ---------------------------------------------------------------------------
# Tool protocol contract
# ---------------------------------------------------------------------------

class ToolContractTests(unittest.TestCase):
    def test_all_four_tools_satisfy_protocol(self):
        from thesis_agent.tools import (
            body_tools,
            heading_tools,
            toc_tools,
            word_postprocess_tools,
        )

        for module in (body_tools, heading_tools, toc_tools, word_postprocess_tools):
            for tool in module.TOOLS:
                with self.subTest(tool=tool.name):
                    self.assertTrue(is_tool(tool))
                    self.assertTrue(tool.name.startswith("tool_"))
                    self.assertIsInstance(tool.input_schema, dict)
                    self.assertIsInstance(tool.requires, list)
                    self.assertIsInstance(tool.idempotent, bool)


# ---------------------------------------------------------------------------
# tool_format_body
# ---------------------------------------------------------------------------

class FormatBodyTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.path = self.tmp / "x.docx"
        _make_docx_with_paragraphs(self.path, ["第一段", "第二段"])
        self.dm = DocumentModel.from_path(str(self.path))

    def tearDown(self):
        self._tmp.cleanup()

    def test_changes_normal_style_and_takes_snapshot(self):
        from thesis_agent.tools.body_tools import FormatBody

        ctx = _make_ctx()
        result = FormatBody().run(
            self.dm,
            {"east_asia_font": "宋体", "size": 12, "line_spacing": 1.5},
            ctx,
        )
        self.assertTrue(result.ok, msg=result.message)
        self.assertIn("Normal", result.changed_styles)
        self.assertEqual(ctx.snapshot_mgr.take_calls, 1)
        self.assertIsNotNone(result.rollback_token)

    def test_swallows_unexpected_exception(self):
        from thesis_agent.tools.body_tools import FormatBody

        ctx = _make_ctx()
        # Force an internal failure by replacing the doc's write context
        # with one that raises. The Tool must convert that into ok=False
        # rather than letting it propagate.
        broken_doc = type("Broken", (), {
            "write": lambda self_: (_ for _ in ()).throw(RuntimeError("boom")),
            "last_changes": None,
        })()
        result = FormatBody().run(broken_doc, {"east_asia_font": "宋体"}, ctx)
        self.assertFalse(result.ok)
        self.assertIn("boom", result.message)


# ---------------------------------------------------------------------------
# tool_assign_heading_styles
# ---------------------------------------------------------------------------

class AssignHeadingStylesTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        self.path = self.tmp / "x.docx"
        _make_docx_with_paragraphs(
            self.path,
            ["第1章 绪论", "1.1 研究背景", "本研究关注..."],
        )
        self.dm = DocumentModel.from_path(str(self.path))

    def tearDown(self):
        self._tmp.cleanup()

    def test_assigns_heading1_to_chapter_paragraph(self):
        from thesis_agent.tools.heading_tools import AssignHeadingStyles

        ctx = _make_ctx(config=_minimal_cfg())
        result = AssignHeadingStyles().run(self.dm, {}, ctx)
        self.assertTrue(result.ok, msg=result.message)
        # Paragraph 0 should now be Heading 1
        self.assertEqual(self.dm.paragraphs()[0].style_name, "Heading 1")


# ---------------------------------------------------------------------------
# tool_word_postprocess (graceful degrade)
# ---------------------------------------------------------------------------

class WordPostprocessTests(unittest.TestCase):
    def test_mode_none_short_circuits(self):
        from thesis_agent.tools.word_postprocess_tools import WordPostprocess

        result = WordPostprocess().run(
            object(),  # docs not used in mode=none
            {"mode": "none"},
            _make_ctx(),
        )
        self.assertTrue(result.ok)
        self.assertIn("none", result.message)


# ---------------------------------------------------------------------------
# tool_insert_toc
# ---------------------------------------------------------------------------

class InsertTocTests(unittest.TestCase):
    def test_declares_heading_dependency(self):
        from thesis_agent.tools.toc_tools import InsertToc

        self.assertEqual(
            InsertToc().requires, ["tool_assign_heading_styles"]
        )


# ---------------------------------------------------------------------------
# Auto-discovery
# ---------------------------------------------------------------------------

class AutoloadTests(unittest.TestCase):
    def test_autoload_registers_four_mvp_tools(self):
        from thesis_agent.tools import registry

        registry.clear()
        registry.autoload()
        names = {t.name for t in registry.all_tools()}
        self.assertIn("tool_format_body", names)
        self.assertIn("tool_assign_heading_styles", names)
        self.assertIn("tool_insert_toc", names)
        self.assertIn("tool_word_postprocess", names)


if __name__ == "__main__":
    unittest.main()
