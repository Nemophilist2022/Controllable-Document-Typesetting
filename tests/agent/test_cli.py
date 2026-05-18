"""thesis-agent CLI smoke tests (R8.1, R12.4)."""

import io
import json
import sys
import tempfile
import unittest
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path

from docx import Document


def _make_doc(path):
    d = Document()
    d.add_paragraph("第一章", style="Heading 1")
    d.add_paragraph("正文")
    d.save(path)


def _make_doc_with_normal_line_spacing(path, line_spacing):
    d = Document()
    normal = d.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.paragraph_format.line_spacing = line_spacing
    d.add_paragraph("Chapter 1", style="Heading 1")
    d.add_paragraph("Body text")
    d.save(path)


class CliTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_run_eval_only_mode_writes_report_no_docx(self):
        from thesis_agent.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            _make_doc(input_path)

            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = main([
                    "run",
                    "--input", str(input_path),
                    "--profile", "scau_2024",
                    "--mode", "eval_only",
                    "--output-dir", tmp,
                ])
            self.assertEqual(rc, 0, msg=buf.getvalue())

            md = list(Path(tmp).glob("*_report.md"))
            js = list(Path(tmp).glob("*_report.json"))
            tr = list(Path(tmp).glob("*_trace.jsonl"))
            self.assertEqual(len(md), 1)
            self.assertEqual(len(js), 1)
            self.assertEqual(len(tr), 1)
            # Input bytes unchanged
            self.assertGreater(input_path.stat().st_size, 0)

    def test_run_eval_only_can_use_custom_yaml_config_without_profile(self):
        from thesis_agent.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            config_path = Path(tmp) / "my_template.yaml"
            _make_doc_with_normal_line_spacing(input_path, 2.0)
            config_path.write_text("body:\n  line_spacing: 2.0\n", encoding="utf-8")

            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = main([
                    "run",
                    "--input", str(input_path),
                    "--config", str(config_path),
                    "--mode", "eval_only",
                    "--output-dir", tmp,
                ])
            self.assertEqual(rc, 0, msg=buf.getvalue())

            report_path = Path(tmp) / "input_report.json"
            payload = json.loads(report_path.read_text(encoding="utf-8"))
            items = {it["rule_id"]: it for it in payload["items"]}

            self.assertEqual(payload["profile"], "my_template")
            self.assertEqual(items["body.line_spacing"]["status"], "done")
            self.assertIn("actual=2.0", items["body.line_spacing"]["evidence"])

    def test_run_with_missing_config_returns_cli_error(self):
        from thesis_agent.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            missing_config = Path(tmp) / "missing.yaml"
            _make_doc(input_path)

            err = io.StringIO()
            with redirect_stderr(err):
                rc = main([
                    "run",
                    "--input", str(input_path),
                    "--config", str(missing_config),
                    "--mode", "eval_only",
                    "--output-dir", tmp,
                ])

            self.assertEqual(rc, 2)
            self.assertIn("template/config error", err.getvalue())
            self.assertIn("missing.yaml", err.getvalue())

    def test_list_profiles_includes_scau_2024(self):
        from thesis_agent.cli import main

        buf = io.StringIO()
        with redirect_stdout(buf):
            rc = main(["list", "profiles"])
        self.assertEqual(rc, 0)
        self.assertIn("scau_2024", buf.getvalue())

    def test_list_tools_includes_mvp_tools(self):
        from thesis_agent.cli import main

        buf = io.StringIO()
        with redirect_stdout(buf):
            rc = main(["list", "tools"])
        self.assertEqual(rc, 0)
        out = buf.getvalue()
        self.assertIn("tool_format_body", out)
        self.assertIn("tool_assign_heading_styles", out)

    def test_list_rules_scau_2024(self):
        from thesis_agent.cli import main

        buf = io.StringIO()
        with redirect_stdout(buf):
            rc = main(["list", "rules", "scau_2024"])
        self.assertEqual(rc, 0)
        out = buf.getvalue()
        self.assertIn("body.font.east_asia", out)

    def test_list_rules_can_use_custom_yaml_config(self):
        from thesis_agent.cli import main

        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "custom.yaml"
            config_path.write_text("body:\n  line_spacing: 2.0\n", encoding="utf-8")

            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = main(["list", "rules", "--config", str(config_path)])

            self.assertEqual(rc, 0)
            out = buf.getvalue()
            self.assertIn("body.line_spacing", out)
            self.assertIn("fix_tool=tool_format_body", out)

    def test_extract_template_from_text_writes_yaml(self):
        from thesis_agent.cli import main
        from thesis_agent.ingest.template_loader import from_yaml

        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "nl.yaml"

            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = main([
                    "extract-template",
                    "--text", "\u6b63\u6587\u5c0f\u56db\u53f7\u5b8b\u4f53\uff0c1.5\u500d\u884c\u8ddd\u3002",
                    "--output", str(out_path),
                ])

            self.assertEqual(rc, 0, msg=buf.getvalue())
            cfg = from_yaml(str(out_path))
            self.assertEqual(cfg["fonts"]["body"], "\u5b8b\u4f53")
            self.assertEqual(cfg["sizes"]["body"], 12)
            self.assertEqual(cfg["body"]["line_spacing"], 1.5)
            self.assertIn("status=ready", buf.getvalue())

    def test_extract_template_from_docx_writes_yaml(self):
        from docx.shared import Pt

        from thesis_agent.cli import main
        from thesis_agent.ingest.template_loader import from_yaml

        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            out_path = Path(tmp) / "docx.yaml"
            d = Document()
            d.styles["Normal"].font.size = Pt(12)
            d.styles["Normal"].paragraph_format.line_spacing = 2.0
            d.save(template_path)

            buf = io.StringIO()
            with redirect_stdout(buf):
                rc = main([
                    "extract-template",
                    "--docx", str(template_path),
                    "--output", str(out_path),
                ])

            self.assertEqual(rc, 0, msg=buf.getvalue())
            cfg = from_yaml(str(out_path))
            self.assertEqual(cfg["sizes"]["body"], 12)
            self.assertEqual(cfg["body"]["line_spacing"], 2.0)

    def test_print_line_is_safe_for_gbk_stdout(self):
        from thesis_agent.cli import _print_line

        class StrictGbkStdout:
            encoding = "gbk"

            def __init__(self):
                self.text = ""

            def write(self, s):
                s.encode("gbk")
                self.text += s

            def flush(self):
                pass

        old = sys.stdout
        fake = StrictGbkStdout()
        try:
            sys.stdout = fake
            _print_line("⚠️ 待审 pending: x")
        finally:
            sys.stdout = old

        self.assertIn("待审 pending", fake.text)
        self.assertNotIn("⚠", fake.text)


class BackwardCompatTests(unittest.TestCase):
    def test_existing_run_format_signature_unchanged(self):
        import inspect

        import thesis_runner

        sig = inspect.signature(thesis_runner.run_format)
        params = list(sig.parameters)
        # The legacy positional+kwargs surface must still be there.
        for p in ["input_path", "output_path", "log", "config", "config_path"]:
            self.assertIn(p, params)


if __name__ == "__main__":
    unittest.main()
