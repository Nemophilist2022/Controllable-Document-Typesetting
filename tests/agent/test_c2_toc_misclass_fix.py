"""C2 fix — TOC entries no longer get promoted to Heading 1.

Two scenarios:

1. ``第一章 绪论 ............ 1`` styled as TOC 1 — must remain TOC 1
2. Same text styled as Normal but with leader dots — must remain Normal

And one regression guard:

3. ``第一章 绪论`` (no leader) styled as Normal — must be promoted to
   Heading 1 (that's the legacy behaviour and the entire reason
   ``auto_assign_heading_styles`` exists).
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from thesis_formatter.headings import (
    _looks_like_toc_entry,
    auto_assign_heading_styles,
)
from thesis_config import DEFAULT_CONFIG


def _ensure_toc1(doc):
    if "TOC 1" not in [s.name for s in doc.styles]:
        doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)


# ---------------------------------------------------------------------------
# Helper unit tests
# ---------------------------------------------------------------------------

class LooksLikeTocEntryUnitTests(unittest.TestCase):
    def setUp(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "x.docx"
            doc = Document()
            _ensure_toc1(doc)
            doc.add_paragraph("dummy", style="TOC 1")
            doc.save(path)
            # Re-open so styles persist; keep a real para object.
            self.doc = Document(path)

    def test_toc1_styled_paragraph(self):
        para = self.doc.paragraphs[0]
        self.assertTrue(_looks_like_toc_entry(para, "第一章 绪论"))

    def test_normal_paragraph_with_leader_dots(self):
        # Build a normal paragraph and feed text with dots.
        doc = Document()
        para = doc.add_paragraph("第一章 绪论 ............ 1")
        self.assertTrue(_looks_like_toc_entry(para, para.text))

    def test_normal_paragraph_without_leader_is_not_toc(self):
        doc = Document()
        para = doc.add_paragraph("第一章 绪论")
        self.assertFalse(_looks_like_toc_entry(para, para.text))


# ---------------------------------------------------------------------------
# Integration: auto_assign_heading_styles must not touch TOC entries
# ---------------------------------------------------------------------------

class AutoAssignTocSafetyTests(unittest.TestCase):
    def test_toc1_styled_chapter_text_stays_toc1(self):
        doc = Document()
        _ensure_toc1(doc)
        # Body chapter — should be promoted to Heading 1.
        doc.add_paragraph("第一章 绪论", style="Normal")
        # TOC entry — must stay TOC 1.
        doc.add_paragraph("第一章 绪论 ............ 1", style="TOC 1")

        auto_assign_heading_styles(doc, DEFAULT_CONFIG)

        # paragraph 0 (chapter) → Heading 1
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        # paragraph 1 (TOC entry) → still TOC 1
        self.assertEqual(doc.paragraphs[1].style.name, "TOC 1")

    def test_normal_paragraph_with_leader_dots_stays_normal(self):
        """Even without a TOC style, a paragraph that looks like a
        TOC entry (leader dots) is left alone."""
        doc = Document()
        # No TOC 1 style registered — a plain Normal paragraph that
        # happens to contain a leader run.
        doc.add_paragraph("第一章 绪论 ............ 1", style="Normal")

        auto_assign_heading_styles(doc, DEFAULT_CONFIG)

        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_real_chapter_heading_still_promoted(self):
        """Regression guard: legacy behaviour for true chapter
        paragraphs must NOT change."""
        doc = Document()
        doc.add_paragraph("第一章 绪论", style="Normal")
        doc.add_paragraph("正文段落。", style="Normal")

        auto_assign_heading_styles(doc, DEFAULT_CONFIG)

        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        # Body paragraph stays untouched.
        self.assertEqual(doc.paragraphs[1].style.name, "Normal")


if __name__ == "__main__":
    unittest.main()
