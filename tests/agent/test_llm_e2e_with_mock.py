"""End-to-end LLM diagnose → fix → re-evaluate loop, using MockLLMClient.

Real LLM calls cost money and require credentials, so we use a mock
that returns canned responses with realistic shape. The mock's
behaviour mirrors what a real OpenAI-compatible response would look
like for our prompts.

The point of this test is to prove the full closed loop works:
    eval finds fail → diagnose returns ToolCall → planner builds new
    plan from ToolCall → act applies it → eval shows pass.

If this passes, swapping ``MockLLMClient`` for ``OpenAICompatibleClient``
is a config-only change (set THESIS_AGENT_LLM_API_KEY).
"""

import json
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _make_violating_doc(path):
    """Normal style with line spacing 2.0 (违 body.line_spacing=1.5)."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 2.0  # ❌ should be 1.5

    doc.add_paragraph("第一章 绪论", style="Heading 1")
    doc.add_paragraph("正文段落。", style="Normal")
    doc.save(path)


class LLMClosedLoopTests(unittest.TestCase):
    def setUp(self):
        from thesis_agent.diagnoser.diagnoser import reset_caches
        from thesis_agent.evaluators import runner
        from thesis_agent.tools import registry

        reset_caches()
        runner.clear_checks()
        registry.clear()

    def test_mock_llm_response_drives_replan_to_pass(self):
        """body.line_spacing fail → mock LLM returns fix_plan calling
        tool_format_body{line_spacing: 1.5} → next iteration passes."""
        from thesis_agent.diagnoser.llm_client import MockLLMClient

        # Inject mock LLM by patching _build_llm_client.
        canned = {
            "body.line_spacing": {
                "rule_id": "body.line_spacing",
                "root_cause": "Normal style line_spacing was set to 2.0",
                "fix_plan": [
                    {
                        "tool": "tool_format_body",
                        "params": {"line_spacing": 1.5},
                        "expected_effect": "Reset to 1.5x line spacing",
                    }
                ],
                "confidence": 0.95,
                "needs_human": False,
                "rationale": "Numeric spacing — straightforward fix.",
            }
        }
        mock_llm = MockLLMClient(canned=canned)

        with tempfile.TemporaryDirectory() as tmp:
            in_path = Path(tmp) / "x.docx"
            _make_violating_doc(in_path)

            from thesis_agent.orchestrator.harness import RunOptions, run

            with mock.patch(
                "thesis_agent.orchestrator.harness._build_llm_client",
                return_value=mock_llm,
            ):
                result = run(
                    input_path=str(in_path),
                    profile="scau_2024",
                    mode="full",
                    options=RunOptions(
                        output_dir=tmp,
                        auto_apply_diagnosis="yes",  # 让 LLM 修复直接执行
                    ),
                )

            self.assertTrue(result.ok)
            # body.line_spacing should now be done after the LLM-driven
            # replan → act cycle.
            report = json.loads(
                Path(result.report_json_path).read_text(encoding="utf-8")
            )
            statuses = {it["rule_id"]: it["status"] for it in report["items"]}
            self.assertEqual(
                statuses.get("body.line_spacing"), "done",
                msg=f"expected body.line_spacing=done after LLM fix, "
                    f"got {statuses.get('body.line_spacing')!r}; "
                    f"summary={report['summary']}",
            )
            # And the trace must show the replan/act cycle actually fired.
            trace_text = Path(result.trace_path).read_text(encoding="utf-8")
            self.assertIn("tool_format_body", trace_text)
            self.assertIn("diagnose", trace_text)

    def test_full_mode_report_records_tool_fix_attempts(self):
        """When full mode fixes body.line_spacing, report.json should
        show the actual tool attempt on that rule item."""
        with tempfile.TemporaryDirectory() as tmp:
            in_path = Path(tmp) / "x.docx"
            _make_violating_doc(in_path)

            from thesis_agent.orchestrator.harness import RunOptions, run

            result = run(
                input_path=str(in_path),
                profile="scau_2024",
                mode="full",
                options=RunOptions(
                    output_dir=tmp,
                    auto_apply_diagnosis="no",
                    llm_disabled=True,
                ),
            )

            report = json.loads(
                Path(result.report_json_path).read_text(encoding="utf-8")
            )
            items = {it["rule_id"]: it for it in report["items"]}
            attempts = items["body.line_spacing"]["fix_attempts"]

            self.assertTrue(attempts)
            self.assertEqual(attempts[0]["tool"], "tool_format_body")
            self.assertTrue(attempts[0]["ok"])
            self.assertEqual(attempts[0]["params"]["line_spacing"], 1.5)
            self.assertGreater(report["meta"]["tool_calls_count"], 0)

    def test_full_mode_initial_plan_uses_custom_config_values(self):
        """Custom YAML values must drive the first ACT phase, not only
        later fallback replans."""
        with tempfile.TemporaryDirectory() as tmp:
            in_path = Path(tmp) / "x.docx"
            config_path = Path(tmp) / "custom.yaml"
            _make_violating_doc(in_path)  # line_spacing is already 2.0
            config_path.write_text("body:\n  line_spacing: 2.0\n", encoding="utf-8")

            from thesis_agent.orchestrator.harness import RunOptions, run

            result = run(
                input_path=str(in_path),
                profile="custom",
                mode="full",
                options=RunOptions(
                    output_dir=tmp,
                    auto_apply_diagnosis="no",
                    llm_disabled=True,
                    config_path=str(config_path),
                ),
            )

            report = json.loads(
                Path(result.report_json_path).read_text(encoding="utf-8")
            )
            items = {it["rule_id"]: it for it in report["items"]}
            attempts = items["body.line_spacing"]["fix_attempts"]

            self.assertTrue(attempts)
            self.assertEqual(attempts[0]["tool"], "tool_format_body")
            self.assertEqual(attempts[0]["params"]["line_spacing"], 2.0)
            self.assertEqual(items["body.line_spacing"]["status"], "done")


if __name__ == "__main__":
    unittest.main()
