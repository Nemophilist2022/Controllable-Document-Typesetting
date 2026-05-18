"""End-to-end Check tests against real DocumentModel fixtures (T8)."""

import copy
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt

from thesis_agent.evaluators import runner
from thesis_agent.evaluators.checks import autoload
from thesis_agent.ingest.document_model import DocumentModel
from thesis_agent.spec.compiler import compile
from thesis_agent.spec.rule_set import Rule, RuleSet
from thesis_config import DEFAULT_CONFIG


def _docx_with_normal_style(path, line_spacing=1.5, east_asia="宋体", size_pt=12):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.size = Pt(size_pt)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        east_asia,
    )
    normal.paragraph_format.line_spacing = line_spacing
    doc.add_paragraph("body").style = normal
    doc.save(path)


def _h1_docx(path, with_heading=True):
    doc = Document()
    if with_heading:
        doc.add_paragraph("第一章", style="Heading 1")
    doc.add_paragraph("正文段落")
    doc.save(path)


def _autoload_once():
    runner.clear_checks()
    autoload()


class CheckBodyTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        _autoload_once()

    def tearDown(self):
        self._tmp.cleanup()

    def test_pass_when_normal_style_matches_expected(self):
        path = self.tmp / "ok.docx"
        _docx_with_normal_style(path, line_spacing=1.5, east_asia="宋体", size_pt=12)
        dm = DocumentModel.from_path(str(path))

        rs = compile(DEFAULT_CONFIG)
        report = runner.evaluate(dm, rs)

        statuses = {r.rule_id: r.status for r in report.results}
        self.assertEqual(statuses.get("body.font.east_asia"), "pass")
        self.assertEqual(statuses.get("body.line_spacing"), "pass")

    def test_fail_when_line_spacing_differs(self):
        path = self.tmp / "bad.docx"
        _docx_with_normal_style(path, line_spacing=2.0)
        dm = DocumentModel.from_path(str(path))

        rs = compile(DEFAULT_CONFIG)
        report = runner.evaluate(dm, rs, only_rule_ids=["body.line_spacing"])
        self.assertEqual(len(report.results), 1)
        self.assertEqual(report.results[0].status, "fail")
        # Evidence should be ≤80 chars and mention both values.
        ev = report.results[0].evidence
        self.assertLessEqual(len(ev), 80)
        self.assertIn("2", ev)


class CheckHeadingTests(unittest.TestCase):
    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self.tmp = Path(self._tmp.name)
        _autoload_once()

    def tearDown(self):
        self._tmp.cleanup()

    def test_pass_when_h1_present(self):
        path = self.tmp / "h1.docx"
        _h1_docx(path, with_heading=True)
        dm = DocumentModel.from_path(str(path))

        rs = compile(DEFAULT_CONFIG)
        report = runner.evaluate(
            dm, rs, only_rule_ids=["heading.h1.style_present"]
        )
        self.assertEqual(report.results[0].status, "pass")

    def test_fail_when_h1_missing(self):
        path = self.tmp / "no_h1.docx"
        _h1_docx(path, with_heading=False)
        dm = DocumentModel.from_path(str(path))

        rs = compile(DEFAULT_CONFIG)
        report = runner.evaluate(
            dm, rs, only_rule_ids=["heading.h1.style_present"]
        )
        self.assertEqual(report.results[0].status, "fail")


class CheckTocTests(unittest.TestCase):
    def setUp(self):
        _autoload_once()

    def test_match_heading_count_passes_when_zero_zero(self):
        # Empty doc: 0 headings, 0 toc entries → match → pass
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "empty.docx"
            doc = Document()
            doc.save(path)
            dm = DocumentModel.from_path(str(path))

            rs = compile(DEFAULT_CONFIG)
            report = runner.evaluate(
                dm, rs, only_rule_ids=["toc.entry_count"]
            )
            self.assertEqual(report.results[0].status, "pass")


if __name__ == "__main__":
    unittest.main()
