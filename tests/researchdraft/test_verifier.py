import tempfile
import unittest
from pathlib import Path


class VerifierAgentTests(unittest.TestCase):
    def test_report_lists_content_and_format_findings(self):
        from researchdraft.agents.verifier_agent import VerifierAgent
        from researchdraft.tools.word_tools import markdown_to_docx

        markdown = "\n".join(
            [
                "# Title",
                "",
                "## 摘要",
                "[待补充：实验指标]",
                "",
                "## 参考文献",
                "[待补充：参考文献]",
            ]
        )

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "paper.docx"
            markdown_to_docx(markdown, docx_path)
            report = VerifierAgent(output_dir=tmp).run(
                draft_markdown=markdown,
                docx_path=str(docx_path),
                trace_entries=[],
            )
            text = Path(report.report_path).read_text("utf-8")

        self.assertIn("[待补充：实验指标]", text)
        self.assertIn("[待补充：参考文献]", text)
        self.assertIn("参考文献区域", text)


class WordOutputEnhancementTests(unittest.TestCase):
    def test_docx_has_centered_title_bold_headings_and_reference_section_break(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from researchdraft.tools.word_tools import markdown_to_docx

        markdown = "\n".join(
            [
                "# Title",
                "",
                "## 摘要",
                "Body text.",
                "",
                "## 参考文献",
                "[待补充：参考文献]",
            ]
        )
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "paper.docx"
            markdown_to_docx(markdown, path)
            from docx import Document

            doc = Document(path)

        self.assertEqual(doc.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        heading_paragraphs = [
            p for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")
        ]
        self.assertTrue(any(p.runs and p.runs[0].bold for p in heading_paragraphs))
        reference = next(p for p in doc.paragraphs if p.text == "参考文献")
        self.assertTrue(reference.paragraph_format.page_break_before)


if __name__ == "__main__":
    unittest.main()
