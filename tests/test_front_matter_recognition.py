import tempfile
import unittest
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from thesis_config import dump_default_config
from thesis_formatter._titles import _detect_front_matter
from thesis_formatter.formatter import apply_format


def _heading_1(doc, text):
    para = doc.add_paragraph(text)
    para.style = doc.styles["Heading 1"]
    return para


def _make_config(mode="auto"):
    cfg = yaml.safe_load(dump_default_config())
    cfg["cover"]["enabled"] = False
    cfg["toc"]["enabled"] = False
    cfg["header_footer"]["enabled"] = False
    cfg["front_matter"]["mode"] = mode
    return cfg


class FrontMatterRecognitionTests(unittest.TestCase):
    def test_detect_front_matter_accepts_standalone_abstract_heading(self):
        doc = Document()
        doc.add_paragraph("Abstract")
        doc.add_paragraph("This is the English abstract body text.")
        _heading_1(doc, "第1章 绪论")

        self.assertTrue(_detect_front_matter(doc, _make_config()))

    def test_auto_mode_formats_standalone_abstract_as_heading(self):
        cfg = _make_config()
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("中文自定义标题")
            doc.add_paragraph("这是一段中文摘要。")
            doc.add_paragraph("关键词: 测试, 前置页")
            doc.add_paragraph("Abstract")
            doc.add_paragraph("This is the English abstract body text.")
            doc.add_paragraph("Key words: alpha, beta")
            _heading_1(doc, "第1章 绪论")
            doc.add_paragraph("正文开始。")
            doc.save(input_path)

            apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)
            abstract_para = next(p for p in out.paragraphs if p.text.strip() == "Abstract")
            english_body_para = next(p for p in out.paragraphs if "English abstract body text" in p.text)

            self.assertEqual(abstract_para.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertNotEqual(abstract_para.style.name, "Heading 1")
            self.assertEqual(english_body_para.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_auto_mode_warns_instead_of_inserting_missing_abstract_title(self):
        cfg = _make_config()
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("原前置标题")
            doc.add_paragraph("这是一段中文摘要。")
            doc.add_paragraph("关键词: 测试, 前置页")
            _heading_1(doc, "第1章 绪论")
            doc.add_paragraph("正文开始。")
            doc.save(input_path)

            warnings = apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)
            first_para = next(p for p in out.paragraphs if p.text.strip())

            self.assertEqual(first_para.text.strip(), "原前置标题")
            self.assertTrue(any("摘要/Abstract" in warning for warning in warnings))

    def test_auto_mode_does_not_warn_when_explicit_cn_abstract_title_exists_later(self):
        cfg = _make_config()
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("封面占位")
            doc.add_paragraph("声明占位")
            doc.add_paragraph("摘        要")
            doc.add_paragraph("这是一段中文摘要。")
            doc.add_paragraph("关键词: 测试, 前置页")
            _heading_1(doc, "第1章 绪论")
            doc.add_paragraph("正文开始。")
            doc.save(input_path)

            warnings = apply_format(str(input_path), str(output_path), config=cfg)

            self.assertFalse(any("摘要/Abstract" in warning for warning in warnings))

    def test_format_mode_warns_instead_of_inserting_missing_abstract_title(self):
        cfg = _make_config(mode="format")
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("原前置标题")
            doc.add_paragraph("这是一段中文摘要。")
            doc.add_paragraph("关键词: 测试, 前置页")
            _heading_1(doc, "第1章 绪论")
            doc.add_paragraph("正文开始。")
            doc.save(input_path)

            warnings = apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)
            first_para = next(p for p in out.paragraphs if p.text.strip())

            self.assertEqual(first_para.text.strip(), "原前置标题")
            self.assertTrue(any("摘要/Abstract" in warning for warning in warnings))

    def test_auto_mode_warns_when_english_abstract_label_missing(self):
        cfg = _make_config()
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("摘要")
            doc.add_paragraph("这是一段中文摘要。")
            doc.add_paragraph("关键词: 测试, 前置页")
            doc.add_paragraph("An English Thesis Title")
            doc.add_paragraph("Alice Example")
            doc.add_paragraph("(Example University, China)")
            doc.add_paragraph("This is the English abstract body text without an explicit label.")
            doc.add_paragraph("Key words: alpha, beta")
            _heading_1(doc, "第1章 绪论")
            doc.add_paragraph("正文开始。")
            doc.save(input_path)

            warnings = apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)
            first_para = next(p for p in out.paragraphs if p.text.strip())

            self.assertFalse(any(p.text.strip() == "Abstract" for p in out.paragraphs))
            self.assertNotEqual(first_para.style.name, "Heading 1")
            self.assertTrue(any("英文Abstract" in warning for warning in warnings))


if __name__ == "__main__":
    unittest.main()
