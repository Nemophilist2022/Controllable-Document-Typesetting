import copy
import tempfile
from pathlib import Path
import unittest

from docx import Document

from thesis_config import DEFAULT_CONFIG
from thesis_formatter._common import get_paragraph_heading_level
from thesis_formatter.formatter import apply_format
from thesis_formatter.headings import auto_assign_heading_styles


class HeadingRecognitionTests(unittest.TestCase):
    def _assign_levels(self, texts):
        cfg = copy.deepcopy(DEFAULT_CONFIG)
        doc = Document()
        paragraphs = [doc.add_paragraph(text) for text in texts]
        changes = auto_assign_heading_styles(doc, cfg)
        levels = [get_paragraph_heading_level(p) for p in paragraphs]
        return changes, levels

    def test_top_aligned_manual_numeric_subheadings_are_recognized(self):
        changes, levels = self._assign_levels([
            "1 绪论",
            "1.1研究背景",
            "1.1.1 技术路线",
            "1.1.1.1实现细节",
        ])

        self.assertEqual(levels, [1, 2, 3, 4])
        self.assertEqual(len(changes), 4)

    def test_long_body_sentences_are_not_misrecognized_as_headings(self):
        changes, levels = self._assign_levels([
            "2025年7月举行的中央城市工作会议明确提出了建设创新、宜居、美丽、韧性、文明、智慧的现代化人民城市这一战略方向。",
            "第一章作为导论部分，重在阐述研究背景与核心问题的提出过程。",
            "1966年，弗里德曼在其著作《区域发展政策》中对此进行了系统性阐述。",
        ])

        self.assertEqual(levels, [None, None, None])
        self.assertEqual(changes, [])

    def test_full_format_normalizes_heading_spacing(self):
        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg["cover"]["enabled"] = False
        cfg["toc"]["enabled"] = False
        cfg["header_footer"]["enabled"] = False

        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            output_path = Path(tmp) / "output.docx"

            doc = Document()
            doc.add_paragraph("1 绪论")
            doc.add_paragraph("1.1研究背景")
            doc.add_paragraph("正文第一段")
            doc.save(input_path)

            apply_format(str(input_path), str(output_path), config=cfg)

            out = Document(output_path)
            meaningful = [p for p in out.paragraphs if p.text.strip()]
            self.assertEqual(meaningful[0].text.strip(), "1  绪论")
            self.assertEqual(meaningful[1].text.strip(), "1.1  研究背景")


if __name__ == "__main__":
    unittest.main()
