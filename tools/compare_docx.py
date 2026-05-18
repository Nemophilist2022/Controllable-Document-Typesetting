"""Semantic-level docx comparator.

Backs requirement R10.2: comparing two docx files at the
"paragraph + section + style" semantic layer, ignoring docx ZIP
metadata such as timestamps, ``rsid*`` tags, ``docPr`` ids, and theme
font lang attributes that fluctuate between builds.

Public API:
    segment_equal(a, b) -> bool
    segment_diff(a, b)  -> list[Diff]

CLI:
    python -m tools.compare_docx <a.docx> <b.docx>
        exit code 0 = equal, 1 = different (diff printed to stdout)
"""

from __future__ import annotations

import sys
from dataclasses import dataclass, field
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Diff datatype
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Diff:
    """A single semantic difference between two docx files.

    Attributes:
        kind: short discriminator, e.g. ``paragraph_text``,
            ``paragraph_count``, ``style_def``, ``section_attr``.
        locator: a dict pinpointing the location, e.g.
            ``{"paragraph_index": 3}`` or ``{"style_name": "Heading 1"}``.
        actual: the value found in *b*.
        expected: the value found in *a* (treated as the reference).
    """

    kind: str
    locator: dict[str, Any] = field(default_factory=dict)
    actual: Any = None
    expected: Any = None


# ---------------------------------------------------------------------------
# Fingerprinting
# ---------------------------------------------------------------------------

# Element / attribute tags that vary between builds and must be ignored.
_IGNORED_ATTRS = {
    qn("w:rsidR"),
    qn("w:rsidRDefault"),
    qn("w:rsidP"),
    qn("w:rsidRPr"),
    qn("w:rsidTr"),
    qn("w:rsidSect"),
    qn("w:paraId"),
    qn("w:textId"),
}


def _run_fingerprint(run) -> dict[str, Any]:
    """Stable representation of a run's text and key formatting."""
    rpr = run._element.find(qn("w:rPr"))
    fp: dict[str, Any] = {"text": run.text or ""}

    if rpr is not None:
        # Bold / italic / underline (presence of the element is what matters).
        for tag, key in (
            (qn("w:b"), "bold"),
            (qn("w:i"), "italic"),
            (qn("w:u"), "underline"),
            (qn("w:strike"), "strike"),
        ):
            el = rpr.find(tag)
            if el is not None:
                val = el.get(qn("w:val"))
                # ``<w:b/>`` (no val) means "on"; ``<w:b w:val="0"/>`` means off.
                if val in (None, "true", "1", "on"):
                    fp[key] = True
                else:
                    fp[key] = False

        # Size / colour.
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            fp["sz"] = sz.get(qn("w:val"))
        color = rpr.find(qn("w:color"))
        if color is not None:
            fp["color"] = color.get(qn("w:val"))

        # Fonts. Capture all four font slots so latin/east-asia changes
        # show up in diffs.
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            fp["fonts"] = {
                "ascii": rfonts.get(qn("w:ascii")),
                "hAnsi": rfonts.get(qn("w:hAnsi")),
                "eastAsia": rfonts.get(qn("w:eastAsia")),
                "cs": rfonts.get(qn("w:cs")),
            }

    return fp


def _para_fingerprint(para) -> dict[str, Any]:
    """Stable representation of a paragraph: style + text + run formatting."""
    style_id = None
    style_name = None
    if para.style is not None:
        style_id = getattr(para.style, "style_id", None)
        style_name = para.style.name

    runs = [_run_fingerprint(r) for r in para.runs]

    # Paragraph-level alignment.
    align = None
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            align = jc.get(qn("w:val"))

    return {
        "style_id": style_id,
        "style_name": style_name,
        "text": para.text or "",
        "align": align,
        "runs": runs,
    }


def _style_fingerprint(style) -> dict[str, Any]:
    """Stable representation of a style definition.

    ``python-docx`` exposes a few specialised style classes
    (e.g. ``_NumberingStyle``) that don't define ``base_style``. Use
    ``getattr`` with a default so those styles still produce a stable
    fingerprint.
    """
    base = getattr(style, "base_style", None)
    return {
        "name": style.name,
        "type": str(style.type) if style.type is not None else None,
        "base": base.name if base is not None else None,
    }


def _section_fingerprint(section) -> dict[str, Any]:
    """Stable representation of section attributes."""

    def _emu(value):
        return None if value is None else int(value)

    return {
        "page_height": _emu(section.page_height),
        "page_width": _emu(section.page_width),
        "left_margin": _emu(section.left_margin),
        "right_margin": _emu(section.right_margin),
        "top_margin": _emu(section.top_margin),
        "bottom_margin": _emu(section.bottom_margin),
        "header_distance": _emu(section.header_distance),
        "footer_distance": _emu(section.footer_distance),
        "gutter": _emu(section.gutter),
        "orientation": str(section.orientation) if section.orientation is not None else None,
        "start_type": str(section.start_type) if section.start_type is not None else None,
    }


# ---------------------------------------------------------------------------
# Diffing
# ---------------------------------------------------------------------------


def _diff_paragraphs(a_paras, b_paras) -> Iterable[Diff]:
    if len(a_paras) != len(b_paras):
        yield Diff(
            kind="paragraph_count",
            locator={},
            actual=len(b_paras),
            expected=len(a_paras),
        )

    for idx, (fa, fb) in enumerate(zip(a_paras, b_paras)):
        if fa["text"] != fb["text"]:
            yield Diff(
                kind="paragraph_text",
                locator={"paragraph_index": idx},
                actual=fb["text"],
                expected=fa["text"],
            )
            continue
        if fa["style_id"] != fb["style_id"] or fa["style_name"] != fb["style_name"]:
            yield Diff(
                kind="paragraph_style",
                locator={"paragraph_index": idx},
                actual=(fb["style_id"], fb["style_name"]),
                expected=(fa["style_id"], fa["style_name"]),
            )
        if fa["align"] != fb["align"]:
            yield Diff(
                kind="paragraph_align",
                locator={"paragraph_index": idx},
                actual=fb["align"],
                expected=fa["align"],
            )
        if fa["runs"] != fb["runs"]:
            yield Diff(
                kind="paragraph_runs",
                locator={"paragraph_index": idx},
                actual=fb["runs"],
                expected=fa["runs"],
            )


def _diff_sections(a_secs, b_secs) -> Iterable[Diff]:
    if len(a_secs) != len(b_secs):
        yield Diff(
            kind="section_count",
            locator={},
            actual=len(b_secs),
            expected=len(a_secs),
        )

    for idx, (fa, fb) in enumerate(zip(a_secs, b_secs)):
        if fa != fb:
            yield Diff(
                kind="section_attr",
                locator={"section_index": idx},
                actual=fb,
                expected=fa,
            )


def _diff_styles(a_styles, b_styles) -> Iterable[Diff]:
    a_by_name = {s["name"]: s for s in a_styles}
    b_by_name = {s["name"]: s for s in b_styles}

    only_in_a = sorted(set(a_by_name) - set(b_by_name))
    only_in_b = sorted(set(b_by_name) - set(a_by_name))

    for name in only_in_a:
        yield Diff(
            kind="style_def",
            locator={"style_name": name},
            actual=None,
            expected=a_by_name[name],
        )
    for name in only_in_b:
        yield Diff(
            kind="style_def",
            locator={"style_name": name},
            actual=b_by_name[name],
            expected=None,
        )

    for name in sorted(set(a_by_name) & set(b_by_name)):
        if a_by_name[name] != b_by_name[name]:
            yield Diff(
                kind="style_def",
                locator={"style_name": name},
                actual=b_by_name[name],
                expected=a_by_name[name],
            )


def _collect(doc) -> tuple[list[dict], list[dict], list[dict]]:
    paras = [_para_fingerprint(p) for p in doc.paragraphs]
    secs = [_section_fingerprint(s) for s in doc.sections]
    styles = [_style_fingerprint(s) for s in doc.styles]
    return paras, secs, styles


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def segment_diff(path_a: str, path_b: str) -> list[Diff]:
    """Return all semantic differences between *path_a* and *path_b*.

    Empty list means the two docx files are semantically equivalent.
    """
    doc_a = Document(path_a)
    doc_b = Document(path_b)

    a_paras, a_secs, a_styles = _collect(doc_a)
    b_paras, b_secs, b_styles = _collect(doc_b)

    diffs: list[Diff] = []
    diffs.extend(_diff_paragraphs(a_paras, b_paras))
    diffs.extend(_diff_sections(a_secs, b_secs))
    diffs.extend(_diff_styles(a_styles, b_styles))
    return diffs


def segment_equal(path_a: str, path_b: str) -> bool:
    """Return True if the two docx files are semantically equivalent."""
    return not segment_diff(path_a, path_b)


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def _main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: python -m tools.compare_docx <a.docx> <b.docx>", file=sys.stderr)
        return 2
    diffs = segment_diff(argv[1], argv[2])
    if not diffs:
        print("OK: semantically equal")
        return 0
    print(f"DIFF: {len(diffs)} difference(s)")
    for d in diffs:
        print(f"  - {d.kind} @ {d.locator}: expected={d.expected!r} actual={d.actual!r}")
    return 1


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv))
