from __future__ import annotations

import copy
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def markdown_to_docx(markdown: str, output_path: str | Path) -> str:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_base_styles(doc)

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bold_runs(p)
        elif line.startswith("## "):
            title = line[3:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            if title in {"参考文献", "References"}:
                p.paragraph_format.page_break_before = True
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            _bold_runs(p)
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            _format_body_paragraph(p)
        else:
            p = doc.add_paragraph(line)
            _format_body_paragraph(p)

    _add_footer_page_number_placeholder(doc)

    doc.save(output)
    return str(output)


def _configure_base_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    for name in ("Title", "Heading 1", "Heading 2"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.bold = True


def _format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def _bold_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.bold = True


def _add_footer_page_number_placeholder(doc) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.text.strip():
            continue
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


class _SnapshotManager:
    def take(self, doc, tool_name=""):
        return f"researchdraft-{tool_name}"


class _NoopTrace:
    def record(self, *, kind: str, payload: dict[str, Any]) -> None:
        return None


def run_word_toolchain(docx_path: str | Path) -> list[dict[str, Any]]:
    from thesis_agent.ingest.document_model import DocumentModel
    from thesis_agent.tools import registry
    from thesis_agent.tools.base import ToolContext
    from thesis_config import DEFAULT_CONFIG

    registry.clear()
    registry.autoload()
    dm = DocumentModel.from_path(str(docx_path))
    cfg = copy.deepcopy(DEFAULT_CONFIG)
    cfg.setdefault("cover", {})["enabled"] = False
    cfg.setdefault("front_matter", {})["mode"] = "skip"

    ctx = ToolContext(
        trace=_NoopTrace(),
        snapshot_mgr=_SnapshotManager(),
        config=cfg,
        runtime={"profile": "researchdraft_mvp", "version": "1"},
    )
    calls = [
        ("tool_assign_heading_styles", {}),
        (
            "tool_format_body",
            {
                "east_asia_font": cfg.get("fonts", {}).get("body", "宋体"),
                "size": cfg.get("sizes", {}).get("body", 12),
                "line_spacing": cfg.get("body", {}).get("line_spacing", 1.5),
                "first_line_indent": cfg.get("body", {}).get("first_line_indent", 24),
                "align": cfg.get("body", {}).get("align", "justify"),
                "latin_font": cfg.get("fonts", {}).get("latin", "Times New Roman"),
            },
        ),
        ("tool_insert_toc", {}),
        ("tool_setup_page_numbers", {}),
        ("tool_format_references", {}),
        ("tool_word_postprocess", {"mode": "full", "docx_path": str(docx_path)}),
    ]

    results: list[dict[str, Any]] = []
    for name, params in calls:
        try:
            tool = registry.get(name)
            if name == "tool_word_postprocess":
                dm.save(str(docx_path))
            result = tool.run(dm, params, ctx)
            results.append(
                {
                    "tool": name,
                    "ok": bool(result.ok),
                    "message": result.message,
                    "warnings": list(result.warnings),
                }
            )
        except Exception as exc:
            results.append(
                {"tool": name, "ok": False, "message": str(exc), "warnings": []}
            )
    dm.save(str(docx_path))
    return results
