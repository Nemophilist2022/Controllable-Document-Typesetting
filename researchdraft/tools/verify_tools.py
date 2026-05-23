from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


@dataclass
class FormatCheck:
    name: str
    passed: bool
    evidence: str


@dataclass
class VerificationResult:
    completed: list[str] = field(default_factory=list)
    missing_items: list[str] = field(default_factory=list)
    confirmation_items: list[str] = field(default_factory=list)
    format_checks: list[FormatCheck] = field(default_factory=list)
    report_path: str = ""

    @property
    def has_format_problem(self) -> bool:
        return any(not check.passed for check in self.format_checks)


def scan_content_markers(markdown: str) -> tuple[list[str], list[str]]:
    missing = sorted(
        set(re.findall(r"(?:【待补充：[^】]+】|\[待补充：[^\]]+\])", markdown))
    )
    confirmations = sorted(
        set(re.findall(r"(?:【待确认：[^】]+】|\[待确认：[^\]]+\])", markdown))
    )
    return missing, confirmations


def check_docx_format(docx_path: str | Path) -> list[FormatCheck]:
    path = Path(docx_path)
    if not path.exists():
        return [FormatCheck("Word 文档", False, f"不存在: {path}")]

    doc = Document(path)
    paragraphs = list(doc.paragraphs)
    heading_count = sum(
        1
        for p in paragraphs
        if p.style is not None and p.style.name.lower().startswith("heading")
    )
    has_toc = any("TOC" in (p.style.name if p.style else "") for p in paragraphs) or any(
        "目录" in p.text or "Table of Contents" in p.text for p in paragraphs
    )
    has_reference = any("参考文献" in p.text or "References" in p.text for p in paragraphs)
    has_body = any(
        p.text.strip()
        and not (p.style and p.style.name.lower().startswith("heading"))
        for p in paragraphs
    )
    has_page_number = any(section.footer.paragraphs for section in doc.sections)

    return [
        FormatCheck("标题层级", heading_count > 0, f"Heading 段落数: {heading_count}"),
        FormatCheck("正文样式", has_body, "存在非标题正文段落" if has_body else "未发现正文段落"),
        FormatCheck("目录", has_toc, "检测到目录区域" if has_toc else "未检测到目录区域"),
        FormatCheck(
            "页码",
            has_page_number,
            "检测到页脚段落" if has_page_number else "未检测到页脚段落",
        ),
        FormatCheck(
            "参考文献区域",
            has_reference,
            "检测到参考文献标题" if has_reference else "未检测到参考文献标题",
        ),
    ]
